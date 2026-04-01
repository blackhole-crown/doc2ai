"""DOCX 文件读取器"""

import os
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None
    print("警告: python-docx 未安装，请运行: pip install python-docx")


class DocxReader:
    """Word 文档读取器"""
    
    def __init__(self, config=None):
        self.config = config or {}
        self.metadata = {
            'paragraphs': 0,
            'tables': 0,
            'truncated': False,
            'line_count': 0,
            'displayed_lines': 0,
            'size': 0,
            'error': None
        }
    
    def read(self, file_path):
        """读取 DOCX 文件"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            self.metadata['error'] = 'file_not_found'
            print(f"文件不存在: {file_path}")
            return None
        
        if Document is None:
            self.metadata['error'] = 'python_docx_not_installed'
            print("错误: python-docx 未安装，请运行: pip install python-docx")
            return None
        
        try:
            # 打开文档
            doc = Document(str(file_path))
            text_parts = []
            
            # 提取段落
            para_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    para_count += 1
            
            # 提取表格（如果配置开启）
            if self.config.get('extract_tables', False) and doc.tables:
                text_parts.append("\n" + "="*50)
                text_parts.append("表格内容：")
                text_parts.append("="*50)
                
                for table_idx, table in enumerate(doc.tables, 1):
                    text_parts.append(f"\n表格 {table_idx}:")
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_parts.append(f"  {' | '.join(row_text)}")
            
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip():
                self.metadata['error'] = 'no_text_content'
                print(f"DOCX 文件没有文本内容: {file_path}")
                return None
            
            # 限制行数
            max_lines = self.config.get('max_lines', 500)
            lines = full_text.split('\n')
            original_lines = len(lines)
            
            if len(lines) > max_lines:
                lines = lines[:max_lines//2] + ['... 省略中间内容 ...'] + lines[-max_lines//2:]
                full_text = '\n'.join(lines)
                truncated = True
            else:
                truncated = False
            
            # 更新元数据
            self.metadata = {
                'paragraphs': para_count,
                'tables': len(doc.tables) if doc.tables else 0,
                'truncated': truncated,
                'line_count': original_lines,
                'displayed_lines': len(lines),
                'size': file_path.stat().st_size,
                'error': None
            }
            
            return full_text
            
        except Exception as e:
            self.metadata['error'] = str(e)
            print(f"读取 DOCX 失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def get_metadata(self):
        """返回元数据"""
        return self.metadata